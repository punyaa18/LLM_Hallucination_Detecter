"""Export paper.tex content into DOCX and PDF formats."""

from __future__ import annotations

import re
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader


def latex_to_plain(text: str) -> str:
    text = text.replace("\\\\", "\n")
    text = re.sub(r"\\texttt\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\cite\{[^}]*\}", "", text)
    text = re.sub(r"\\label\{[^}]*\}", "", text)
    text = re.sub(r"\\ref\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\url\{([^}]*)\}", r"\1", text)
    text = text.replace("~", " ")
    text = text.replace("$", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def parse_paper(tex_path: Path):
    raw = tex_path.read_text(encoding="utf-8")

    title_match = re.search(r"\\title\{(.+?)\}", raw, re.DOTALL)
    author_match = re.search(r"\\author\{(.+?)\}", raw, re.DOTALL)
    abstract_match = re.search(r"\\begin\{abstract\}(.+?)\\end\{abstract\}", raw, re.DOTALL)

    title = latex_to_plain(title_match.group(1)) if title_match else ""
    author = latex_to_plain(author_match.group(1)) if author_match else ""
    abstract = latex_to_plain(abstract_match.group(1)) if abstract_match else ""

    sections = []
    figures = []

    figure_blocks = re.findall(r"\\begin\{figure\}\[!t\](.*?)\\end\{figure\}", raw, re.DOTALL)
    for block in figure_blocks:
        path_match = re.search(r"\\includegraphics\[.*?\]\{([^}]*)\}", block)
        cap_match = re.search(r"\\caption\{([^}]*)\}", block, re.DOTALL)
        if path_match:
            figures.append(
                {
                    "path": path_match.group(1).strip(),
                    "caption": latex_to_plain(cap_match.group(1)) if cap_match else "Figure",
                }
            )
    pattern = re.compile(r"\\section\*?\{([^}]*)\}")
    matches = list(pattern.finditer(raw))
    for idx, match in enumerate(matches):
        name = latex_to_plain(match.group(1))
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else raw.find("\\begin{thebibliography}")
        if end == -1:
            end = len(raw)
        body = raw[start:end]

        body = re.sub(r"\\begin\{enumerate\}", "", body)
        body = re.sub(r"\\end\{enumerate\}", "", body)
        body = re.sub(r"\\begin\{itemize\}", "", body)
        body = re.sub(r"\\end\{itemize\}", "", body)
        body = re.sub(r"\\item", "-", body)
        body = re.sub(r"\\begin\{figure\}\[!t\].*?\\end\{figure\}", "", body, flags=re.DOTALL)
        body = re.sub(r"\\begin\{table\}\[!t\].*?\\end\{table\}", "", body, flags=re.DOTALL)
        body = latex_to_plain(body)
        sections.append((name, body))

    refs = []
    bib = re.search(r"\\begin\{thebibliography\}\{00\}(.+?)\\end\{thebibliography\}", raw, re.DOTALL)
    if bib:
        entries = re.findall(r"\\bibitem\{[^}]*\}\s*(.+?)(?=\\bibitem|$)", bib.group(1), re.DOTALL)
        refs = [latex_to_plain(item) for item in entries]

    return title, author, abstract, sections, refs, figures


def export_docx(out_path: Path, title: str, author: str, abstract: str, sections, refs, figures, base_dir: Path):
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(author)

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(abstract)

    for sec_title, sec_body in sections:
        doc.add_heading(sec_title, level=1)
        if sec_body:
            doc.add_paragraph(sec_body)

    if figures:
        doc.add_heading("Figures", level=1)
        for idx, fig in enumerate(figures, 1):
            fig_path = (base_dir / fig["path"]).resolve()
            if fig_path.exists():
                doc.add_picture(str(fig_path), width=Inches(6.1))
                doc.add_paragraph(f"Fig. {idx}. {fig['caption']}")

    if refs:
        doc.add_heading("References", level=1)
        for idx, ref in enumerate(refs, 1):
            doc.add_paragraph(f"[{idx}] {ref}")

    doc.save(str(out_path))


def export_pdf(out_path: Path, title: str, author: str, abstract: str, sections, refs, figures, base_dir: Path):
    c = canvas.Canvas(str(out_path), pagesize=A4)
    width, height = A4
    margin = 50
    y = height - margin

    def write_line(line: str, font="Times-Roman", size=11, gap=15):
        nonlocal y
        if y < margin:
            c.showPage()
            y = height - margin
        c.setFont(font, size)
        c.drawString(margin, y, line)
        y -= gap

    write_line(title, font="Times-Bold", size=14, gap=18)
    write_line(author, font="Times-Roman", size=11, gap=18)

    write_line("Abstract", font="Times-Bold", size=12)
    for ln in wrap(abstract, 105):
        write_line(ln)
    y -= 5

    for sec_title, sec_body in sections:
        write_line(sec_title, font="Times-Bold", size=12)
        for para in sec_body.split(" -"):
            para = para.strip()
            if not para:
                continue
            for ln in wrap(para, 105):
                write_line(ln)
            y -= 2
        y -= 4

    if figures:
        write_line("Figures", font="Times-Bold", size=12)
        for idx, fig in enumerate(figures, 1):
            fig_path = (base_dir / fig["path"]).resolve()
            if not fig_path.exists():
                continue

            image = ImageReader(str(fig_path))
            img_w, img_h = image.getSize()
            max_w = width - 2 * margin
            max_h = 220
            scale = min(max_w / img_w, max_h / img_h)
            draw_w = img_w * scale
            draw_h = img_h * scale

            if y - draw_h - 30 < margin:
                c.showPage()
                y = height - margin

            c.drawImage(str(fig_path), margin, y - draw_h, width=draw_w, height=draw_h, preserveAspectRatio=True, mask='auto')
            y -= draw_h + 14
            for ln in wrap(f"Fig. {idx}. {fig['caption']}", 105):
                write_line(ln, size=10, gap=13)
            y -= 5

    if refs:
        write_line("References", font="Times-Bold", size=12)
        for idx, ref in enumerate(refs, 1):
            for ln in wrap(f"[{idx}] {ref}", 105):
                write_line(ln)

    c.save()


def main():
    base = Path(__file__).resolve().parent
    tex_path = base / "paper.tex"
    docx_path = base / "paper.docx"
    pdf_path = base / "paper.pdf"

    title, author, abstract, sections, refs, figures = parse_paper(tex_path)
    export_docx(docx_path, title, author, abstract, sections, refs, figures, base)
    export_pdf(pdf_path, title, author, abstract, sections, refs, figures, base)

    print(f"Created {docx_path}")
    print(f"Created {pdf_path}")


if __name__ == "__main__":
    main()
